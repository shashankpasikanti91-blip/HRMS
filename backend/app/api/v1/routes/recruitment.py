from __future__ import annotations

from typing import Optional

from fastapi import APIRouter, Depends, File, Form, Query, UploadFile
from pydantic import BaseModel
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.database import get_db
from app.core.deps import get_current_user, require_hr_or_above, require_recruiter_or_above
from app.core.pagination import PaginationParams, Page
from app.models.user import User
from app.schemas.recruitment import (
    JobPostingCreate,
    JobPostingUpdate,
    JobPostingResponse,
    CandidateCreate,
    CandidateUpdate,
    CandidateResponse,
    ApplicationCreate,
    ApplicationStageUpdate,
    ApplicationResponse,
    InterviewCreate,
    InterviewUpdate,
    InterviewResponse,
    OfferCreate,
    OfferUpdate,
    OfferResponse,
    OfferStatusUpdate,
)
from app.services.recruitment_service import RecruitmentService, InterviewService, OfferService

job_router = APIRouter(prefix="/jobs", tags=["Job Postings"])
candidate_router = APIRouter(prefix="/candidates", tags=["Candidates"])
application_router = APIRouter(prefix="/applications", tags=["Applications"])
interview_router = APIRouter(prefix="/interviews", tags=["Interviews"])
offer_router = APIRouter(prefix="/offers", tags=["Offers"])
ai_recruitment_router = APIRouter(prefix="/recruitment-ai", tags=["AI Recruitment"])


# ── Job Postings ───────────────────────────────────────────────────────────

@job_router.post("", response_model=JobPostingResponse, status_code=201)
async def create_job(
    data: JobPostingCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_recruiter_or_above()),
):
    svc = RecruitmentService(db)
    job = await svc.create_job(data, current_user.company_id, current_user.id)
    return JobPostingResponse.model_validate(job)


@job_router.get("", response_model=Page[JobPostingResponse])
async def list_jobs(
    params: PaginationParams = Depends(),
    status: Optional[str] = Query(None),
    department_id: Optional[str] = Query(None),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    svc = RecruitmentService(db)
    jobs, total = await svc.list_jobs(
        current_user.company_id, params, status=status
    )
    return Page.create([JobPostingResponse.model_validate(j) for j in jobs], total, params)


@job_router.get("/{business_id}", response_model=JobPostingResponse)
async def get_job(
    business_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    svc = RecruitmentService(db)
    job = await svc.get_job(business_id, current_user.company_id)
    return JobPostingResponse.model_validate(job)


@job_router.put("/{business_id}", response_model=JobPostingResponse)
async def update_job(
    business_id: str,
    data: JobPostingUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_recruiter_or_above()),
):
    svc = RecruitmentService(db)
    job = await svc.update_job(business_id, data, current_user.company_id, current_user.id)
    return JobPostingResponse.model_validate(job)


@job_router.delete("/{business_id}", status_code=204)
async def delete_job(
    business_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_hr_or_above()),
):
    svc = RecruitmentService(db)
    await svc.delete_job(business_id, current_user.company_id)


# ── Candidates ─────────────────────────────────────────────────────────────

@candidate_router.post("", response_model=CandidateResponse, status_code=201)
async def create_candidate(
    data: CandidateCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_recruiter_or_above()),
):
    svc = RecruitmentService(db)
    candidate = await svc.create_candidate(data, current_user.company_id, current_user.id)
    return CandidateResponse.model_validate(candidate)


@candidate_router.get("", response_model=Page[CandidateResponse])
async def list_candidates(
    params: PaginationParams = Depends(),
    search: Optional[str] = Query(None),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    if search:
        params.q = search
    svc = RecruitmentService(db)
    candidates, total = await svc.list_candidates(current_user.company_id, params)
    return Page.create([CandidateResponse.model_validate(c) for c in candidates], total, params)


@candidate_router.get("/{business_id}", response_model=CandidateResponse)
async def get_candidate(
    business_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    svc = RecruitmentService(db)
    candidate = await svc.get_candidate(business_id, current_user.company_id)
    return CandidateResponse.model_validate(candidate)


@candidate_router.put("/{business_id}", response_model=CandidateResponse)
async def update_candidate(
    business_id: str,
    data: CandidateUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_recruiter_or_above()),
):
    svc = RecruitmentService(db)
    candidate = await svc.update_candidate(business_id, data, current_user.company_id, current_user.id)
    return CandidateResponse.model_validate(candidate)


@candidate_router.post("/{business_id}/resume", response_model=CandidateResponse)
async def upload_candidate_resume(
    business_id: str,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_recruiter_or_above()),
):
    """Upload a resume file for a candidate."""
    from app.integrations.storage_service import StorageService
    svc = RecruitmentService(db)
    candidate = await svc.get_candidate(business_id, current_user.company_id)
    url = await StorageService.upload(file, folder="resumes", company_id=current_user.company_id)
    candidate.resume_url = url
    candidate.updated_by = current_user.id
    await db.flush()
    await db.refresh(candidate)
    return CandidateResponse.model_validate(candidate)


# ── Applications ───────────────────────────────────────────────────────────

@application_router.post("", response_model=ApplicationResponse, status_code=201)
async def create_application(
    data: ApplicationCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_recruiter_or_above()),
):
    svc = RecruitmentService(db)
    app_obj = await svc.create_application(data, current_user.company_id, current_user.id)
    return ApplicationResponse.model_validate(app_obj)


@application_router.get("", response_model=Page[ApplicationResponse])
async def list_applications(
    params: PaginationParams = Depends(),
    job_id: Optional[str] = Query(None),
    candidate_id: Optional[str] = Query(None),
    stage: Optional[str] = Query(None),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    svc = RecruitmentService(db)
    apps, total = await svc.list_applications(
        current_user.company_id, params, job_id=job_id, candidate_id=candidate_id, stage=stage
    )
    return Page.create([ApplicationResponse.model_validate(a) for a in apps], total, params)


@application_router.get("/{business_id}", response_model=ApplicationResponse)
async def get_application(
    business_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    svc = RecruitmentService(db)
    app_obj = await svc.get_application(business_id, current_user.company_id)
    return ApplicationResponse.model_validate(app_obj)


@application_router.put("/{business_id}/stage", response_model=ApplicationResponse)
async def update_application_stage(
    business_id: str,
    data: ApplicationStageUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_recruiter_or_above()),
):
    """Move application to a new recruitment stage."""
    svc = RecruitmentService(db)
    app_obj = await svc.update_stage(business_id, data, current_user.company_id, current_user.id)
    return ApplicationResponse.model_validate(app_obj)


# ── Interviews ─────────────────────────────────────────────────────────────

@interview_router.post("", response_model=InterviewResponse, status_code=201)
async def create_interview(
    data: InterviewCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_recruiter_or_above()),
):
    svc = InterviewService(db)
    interview = await svc.create(data, current_user.company_id, current_user.id)
    return InterviewResponse.model_validate(interview)


@interview_router.get("", response_model=Page[InterviewResponse])
async def list_interviews(
    params: PaginationParams = Depends(),
    application_id: Optional[str] = Query(None),
    status: Optional[str] = Query(None),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    svc = InterviewService(db)
    interviews, total = await svc.list(
        current_user.company_id, params, application_id=application_id, status=status
    )
    return Page.create([InterviewResponse.model_validate(i) for i in interviews], total, params)


@interview_router.get("/{business_id}", response_model=InterviewResponse)
async def get_interview(
    business_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    svc = InterviewService(db)
    interview = await svc.get(business_id, current_user.company_id)
    return InterviewResponse.model_validate(interview)


@interview_router.put("/{business_id}", response_model=InterviewResponse)
async def update_interview(
    business_id: str,
    data: InterviewUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_recruiter_or_above()),
):
    svc = InterviewService(db)
    interview = await svc.update(business_id, data, current_user.company_id, current_user.id)
    return InterviewResponse.model_validate(interview)


@interview_router.delete("/{business_id}", status_code=204)
async def delete_interview(
    business_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_recruiter_or_above()),
):
    svc = InterviewService(db)
    await svc.delete(business_id, current_user.company_id)


# ── Offers ─────────────────────────────────────────────────────────────────

@offer_router.post("", response_model=OfferResponse, status_code=201)
async def create_offer(
    data: OfferCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_hr_or_above()),
):
    svc = OfferService(db)
    offer = await svc.create(data, current_user.company_id, current_user.id)
    return OfferResponse.model_validate(offer)


@offer_router.get("", response_model=Page[OfferResponse])
async def list_offers(
    params: PaginationParams = Depends(),
    application_id: Optional[str] = Query(None),
    status: Optional[str] = Query(None),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    svc = OfferService(db)
    offers, total = await svc.list(
        current_user.company_id, params, application_id=application_id, status=status
    )
    return Page.create([OfferResponse.model_validate(o) for o in offers], total, params)


@offer_router.get("/{business_id}", response_model=OfferResponse)
async def get_offer(
    business_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    svc = OfferService(db)
    offer = await svc.get(business_id, current_user.company_id)
    return OfferResponse.model_validate(offer)


@offer_router.put("/{business_id}", response_model=OfferResponse)
async def update_offer(
    business_id: str,
    data: OfferUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_hr_or_above()),
):
    svc = OfferService(db)
    offer = await svc.update(business_id, data, current_user.company_id, current_user.id)
    return OfferResponse.model_validate(offer)


@offer_router.put("/{business_id}/status", response_model=OfferResponse)
async def update_offer_status(
    business_id: str,
    data: OfferStatusUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Candidate accepts or declines an offer."""
    svc = OfferService(db)
    offer = await svc.update_status(business_id, data, current_user.company_id, current_user.id)
    return OfferResponse.model_validate(offer)


# ── AI Recruitment ─────────────────────────────────────────────────────────


class AIScreenRequest(BaseModel):
    job_description: str
    resume_text: str


class AIJobPostRequest(BaseModel):
    job_description: str


class AIScreenFileRequest(BaseModel):
    job_posting_id: str
    candidate_id: str


@ai_recruitment_router.post("/screen")
async def ai_screen_candidate(
    data: AIScreenRequest,
    current_user: User = Depends(require_recruiter_or_above()),
):
    """Screen a candidate resume against a job description using AI."""
    from app.core.config import get_settings
    from app.services.ai_recruitment_service import screen_candidate

    settings = get_settings()
    result = await screen_candidate(data.job_description, data.resume_text, settings)
    return result


@ai_recruitment_router.post("/screen-file")
async def ai_screen_candidate_file(
    job_description: str = Form(...),
    resume: UploadFile = File(...),
    current_user: User = Depends(require_recruiter_or_above()),
):
    """Screen a candidate using uploaded resume file against a job description."""
    from app.core.config import get_settings
    from app.services.ai_recruitment_service import screen_candidate

    settings = get_settings()

    # Read resume content
    content = await resume.read()
    resume_text = ""

    filename = (resume.filename or "").lower()
    if filename.endswith(".pdf"):
        try:
            import io
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            resume_text = "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception:
            resume_text = content.decode("utf-8", errors="ignore")
    elif filename.endswith((".doc", ".docx")):
        try:
            import io
            import docx
            doc = docx.Document(io.BytesIO(content))
            resume_text = "\n".join(para.text for para in doc.paragraphs)
        except Exception:
            resume_text = content.decode("utf-8", errors="ignore")
    else:
        resume_text = content.decode("utf-8", errors="ignore")

    if not resume_text.strip():
        return {"error": "Could not extract text from the uploaded file"}

    result = await screen_candidate(job_description, resume_text, settings)
    return result


@ai_recruitment_router.post("/screen-application/{application_business_id}")
async def ai_screen_application(
    application_business_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_recruiter_or_above()),
):
    """Screen a candidate application using stored resume and job description."""
    from app.core.config import get_settings
    from app.services.ai_recruitment_service import screen_candidate
    from app.models.recruitment import JobPosting, Candidate, Application
    from sqlalchemy import select

    settings = get_settings()
    svc = RecruitmentService(db)

    app_obj = await svc.get_application(application_business_id, current_user.company_id)

    # Get job posting
    job_result = await db.execute(
        select(JobPosting).where(JobPosting.id == app_obj.job_posting_id)
    )
    job = job_result.scalar_one_or_none()
    if not job:
        return {"error": "Job posting not found"}

    # Get candidate
    cand_result = await db.execute(
        select(Candidate).where(Candidate.id == app_obj.candidate_id)
    )
    candidate = cand_result.scalar_one_or_none()
    if not candidate:
        return {"error": "Candidate not found"}

    job_desc = f"Title: {job.title}\nDescription: {job.description or ''}\nRequirements: {job.requirements or ''}\nLocation: {job.location or ''}\nType: {job.employment_type or ''}"

    resume_text = ""
    if candidate.resume_url:
        # Try to read resume from local storage
        import os
        local_path = candidate.resume_url.lstrip("/")
        if os.path.exists(local_path):
            try:
                if local_path.endswith(".pdf"):
                    import PyPDF2
                    with open(local_path, "rb") as f:
                        reader = PyPDF2.PdfReader(f)
                        resume_text = "\n".join(page.extract_text() or "" for page in reader.pages)
                elif local_path.endswith((".doc", ".docx")):
                    import docx
                    doc = docx.Document(local_path)
                    resume_text = "\n".join(para.text for para in doc.paragraphs)
                else:
                    with open(local_path, "r", errors="ignore") as f:
                        resume_text = f.read()
            except Exception:
                pass

    if not resume_text.strip():
        # Build resume text from candidate fields
        resume_text = f"""Name: {candidate.full_name}
Email: {candidate.email}
Phone: {candidate.phone or 'N/A'}
Current Company: {candidate.current_company or 'N/A'}
Current Role: {candidate.current_role or 'N/A'}
Experience: {candidate.years_of_experience or 'N/A'} years
Location: {candidate.current_location or 'N/A'}"""

    result = await screen_candidate(job_desc, resume_text, settings)

    # Update application with AI screening results
    if result.get("score"):
        app_obj.ai_screening_score = float(result["score"])
        app_obj.ai_screening_status = "completed"
        if result.get("evaluation", {}).get("justification"):
            app_obj.ai_recommendation = result["evaluation"]["justification"]
        if result.get("evaluation", {}).get("high_match_skills"):
            app_obj.ai_matched_skills = result["evaluation"]["high_match_skills"]
        if result.get("evaluation", {}).get("low_or_missing_match_skills"):
            app_obj.ai_missing_skills = result["evaluation"]["low_or_missing_match_skills"]
        # Also update candidate ai_score
        candidate.ai_score = float(result["score"])
        candidate.ai_summary = result.get("decision", "")
        await db.flush()

    return result


@ai_recruitment_router.post("/generate-job-posts")
async def ai_generate_job_posts(
    data: AIJobPostRequest,
    current_user: User = Depends(require_recruiter_or_above()),
):
    """Generate multi-platform job posts from a job description using AI."""
    from app.core.config import get_settings
    from app.services.ai_recruitment_service import generate_job_posts

    settings = get_settings()
    result = await generate_job_posts(data.job_description, settings)
    return result


@ai_recruitment_router.post("/generate-job-posts-file")
async def ai_generate_job_posts_file(
    jd_file: UploadFile = File(...),
    current_user: User = Depends(require_recruiter_or_above()),
):
    """Generate multi-platform job posts from an uploaded JD file."""
    from app.core.config import get_settings
    from app.services.ai_recruitment_service import generate_job_posts

    settings = get_settings()

    content = await jd_file.read()
    jd_text = ""

    filename = (jd_file.filename or "").lower()
    if filename.endswith(".pdf"):
        try:
            import io
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            jd_text = "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception:
            jd_text = content.decode("utf-8", errors="ignore")
    elif filename.endswith((".doc", ".docx")):
        try:
            import io
            import docx
            doc = docx.Document(io.BytesIO(content))
            jd_text = "\n".join(para.text for para in doc.paragraphs)
        except Exception:
            jd_text = content.decode("utf-8", errors="ignore")
    else:
        jd_text = content.decode("utf-8", errors="ignore")

    if not jd_text.strip():
        return {"error": "Could not extract text from the uploaded file"}

    result = await generate_job_posts(jd_text, settings)
    return result
